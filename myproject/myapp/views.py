from django.shortcuts import render, redirect
import os
from django.conf import settings
from django.http import HttpResponse
import io
from dotenv import load_dotenv

load_dotenv()

import fitz 
from openai import OpenAI
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from docx.oxml import OxmlElement, ns

# from django.http import HttpResponse

def set_column_width(column, width):
    width_twips = int(width.cm * 567)  
    for cell in column.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        for element in tcPr.findall("w:tcW", ns.nsmap):
            tcPr.remove(element)

        tcW = OxmlElement('w:tcW')
        tcW.set(ns.qn('w:w'), str(width_twips))  
        tcW.set(ns.qn('w:type'), 'dxa')  
        
        tcPr.append(tcW)

def pdf_to_text(uploaded_file):
    """Extract text from uploaded PDF file"""
    if uploaded_file is None:
        return ""
    text = ""
    try:
        pdf_bytes = b"".join(chunk for chunk in uploaded_file.chunks())
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            text = "\n".join([page.get_text() for page in doc])
    except Exception as e:
        print(f"Error processing PDF: {e}")
        text = ""
    return text

client = OpenAI(
  base_url="https://openrouter.ai/api/v1",
  api_key=os.getenv("OPENROUTER_API_KEY"),
)



# Create your views here.

def index(request):
    if request.method == "POST":
        return redirect("login")
    return render(request, "index.html")

def login(request):
    return render(request,"login.html")

def dashboard(request):
    context = {}
    if request.method == "POST":
        context = {
            "username" : "Puneeth R V",
            "email" : request.POST.get("email", "")
        }
        return redirect("QPMaker")
    return render(request, "dashboard.html", context)

def QPMaker(request):
    context = {}
    if request.method == "POST":
        context = {
            "subjectMarks" : request.POST.get("subjectMarks", ""),
            "subject" : request.POST.get("subject", ""),
            "diff" : request.POST.get("diff", ""),
            "quesNo" : request.POST.get("quesNo", ""),
            "time" : request.POST.get("time", ""),
            "file" : request.FILES.get("file", None),
            "choice" : request.POST.get("choice")

        }

        ques_marks = context["subjectMarks"]
        ques_subject = context["subject"]
        ques_type = context["diff"]
        ques_num = context["quesNo"]
        ques_time = context["time"]
        pdf_text = pdf_to_text(context["file"])

        content_prompt = ""
        if context["choice"] == "only":
            content_prompt = f"You MUST generate questions STRICTLY and EXCLUSIVELY from the following content. Do NOT use any outside knowledge. Every question must be directly based on this text:\n\n{pdf_text}"
        elif context["choice"] == "also":
            content_prompt = f"Consider also this additional content: {pdf_text}"

        document = Document()

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  

        paragraph2 = document.add_paragraph("[Enter your faculty] \n [Enter name of the test]")
        paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.add_paragraph("_________________________________________________________________________________________________________") 
        document.add_paragraph("Department                                    :           [Enter your department] \nProgramme                                    :           [Enter your program] \nSem/batch                                      :           [] \nDate of test                                     :           [Enter date] \nCourse code                                   :           [Enter Course code] \nCourse Title                                   :           [Enter course title]")
        document.add_paragraph("_________________________________________________________________________________________________________") 

        document.add_paragraph(f"Maximum Duration : {ques_time}                                                                                     Maximum Marks : {ques_marks}")
        document.add_paragraph("_________________________________________________________________________________________________________") 


        # Add an image (Ensure the path is correct)
        run = paragraph.add_run()
        image_path = os.path.join(settings.BASE_DIR, 'static', 'ram.png')
        run.add_picture(image_path, width=Cm(5))

        # OpenAI request to generate questions
        completion = client.chat.completions.create(
        model="deepseek/deepseek-v3.2",
        messages=[
            {
            "role": "user",
            "content": f"""Create a question paper of marks {ques_marks} for the subject {ques_subject} with {ques_type} level questions. 
            There should be {ques_num} questions. {content_prompt}

            Things to keep in mind:
            - Dont number the question
            - dont use latex
            - Mention the marks of each question before the question in "()" in the same line
            - End every question with (OoO)
            
            Do not include:
            - Bold text
            - Your existence in the sentences (Do not refer to yourself as an AI)
            - Extraneous information except the given content at the end by the user (Only generate questions)
            

            

            """
            }
        ]
        )


        questions = completion.choices[0].message.content.strip().split("(OoO)")

        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                    
        header_cells[0].text = 'No.'
        header_cells[1].text = '                                                                 Questions'
        header_cells[2].text = "Marks"


        for idx, question in enumerate(questions, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)  
            row_cells[1].text = question.strip()  
            lastidx = row_cells[1].text.find(")")
            row_cells[2].text = row_cells[1].text[1:lastidx]
            row_cells[1].text = row_cells[1].text[row_cells[1].text.find(")") + 2:]

        set_column_width(table.columns[0], Cm(1))  
        set_column_width(table.columns[1], Cm(17))  
        set_column_width(table.columns[2], Cm(1))

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        response['Content-Disposition'] = f'attachment; filename="{ques_subject}.docx"'
        return response



    return render(request, "QPMaker.html", context)
